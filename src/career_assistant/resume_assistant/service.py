"""简历助手的分析、生成和显式保存服务。"""

from __future__ import annotations

import json
import re
import shutil
import zipfile
from difflib import SequenceMatcher
from dataclasses import asdict
from datetime import datetime, timezone
from pathlib import Path
from xml.sax.saxutils import escape as xml_escape
from uuid import UUID, uuid4

from src.career_assistant.attachments import AttachmentParser
from src.career_assistant.contracts import (
    AttachmentDescriptor,
    ModelCapability,
    ModelSelectionMode,
    ModelSelectionRequest,
)
from src.career_assistant.model_clients import ChatMessage, OpenAICompatibleChatClient
from src.career_assistant.model_gateway import ModelGateway, ModelReadiness, ModelResolution
from src.career_assistant.persistence.model_profile_repository import CareerModelProfileRepository
from src.career_assistant.resume_assistant.models import (
    ResumeAnalysisResult,
    ResumeOptimizationRecord,
    ResumeSuggestion,
)
from src.career_assistant.resume_assistant.repository import ResumeOptimizationRepository
from src.career_assistant.legacy_office import (
    LegacyOfficeConversionError,
    LegacyOfficeConverter,
)
from src.observability.langsmith_runtime import trace_operation


class ResumeOptimizationService:
    """串联既有解析器和模型库，但把历史写入留给用户显式确认。"""

    def __init__(
        self,
        repository: ResumeOptimizationRepository,
        model_repository: CareerModelProfileRepository,
        model_gateway: ModelGateway,
        chat_client: OpenAICompatibleChatClient,
        attachment_parser: AttachmentParser,
        storage_root: Path,
        legacy_office_converter: LegacyOfficeConverter | None = None,
    ) -> None:
        self._repository = repository
        self._model_repository = model_repository
        self._model_gateway = model_gateway
        self._chat_client = chat_client
        self._attachment_parser = attachment_parser
        self._legacy_office_converter = legacy_office_converter
        self._storage_root = storage_root.resolve()

    def analyze(
        self,
        organization_id: UUID,
        resume: AttachmentDescriptor,
        job_description: str,
        job_screenshots: tuple[AttachmentDescriptor, ...],
        extra_prompt: str,
        model_profile_id: UUID | None,
    ) -> ResumeAnalysisResult:
        """解析材料并输出结构化建议，不写历史。"""

        return trace_operation(
            run_name="resume_assistant.analyze",
            run_type="chain",
            inputs={"screenshot_count": len(job_screenshots), "has_extra_prompt": bool(extra_prompt)},
            metadata={"component": "resume_assistant"},
            execute=lambda: self._analyze(
                organization_id,
                resume,
                job_description,
                job_screenshots,
                extra_prompt,
                model_profile_id,
            ),
            summarize=lambda result: {
                "suggestion_count": len(result.suggestions),
                "original_characters": len(result.original_markdown),
            },
            tags=("career", "resume-assistant"),
        )

    def _analyze(
        self,
        organization_id: UUID,
        resume: AttachmentDescriptor,
        job_description: str,
        job_screenshots: tuple[AttachmentDescriptor, ...],
        extra_prompt: str,
        model_profile_id: UUID | None,
    ) -> ResumeAnalysisResult:
        parsed_resume = self._attachment_parser.parse(resume)
        original_markdown = parsed_resume.extracted_text.strip()
        if not original_markdown:
            raise ValueError("未能从简历中解析出可读文字")
        screenshot_texts = []
        for index, descriptor in enumerate(job_screenshots, start=1):
            parsed = self._attachment_parser.parse(descriptor)
            if parsed.extracted_text.strip():
                screenshot_texts.append(f"岗位截图 {index}：\n{parsed.extracted_text.strip()}")
        combined_job = "\n\n".join(
            part for part in (job_description.strip(), *screenshot_texts) if part
        )
        if not combined_job:
            raise ValueError("请填写岗位介绍或上传岗位截图")
        resolution = self._resolve_model(organization_id, model_profile_id)
        prompt = self._analysis_prompt(original_markdown, combined_job, extra_prompt)
        raw = self._chat_client.complete(
            resolution.profile,
            resolution.credential_env_name,
            [
                ChatMessage(role="system", content="你是严谨的技术岗位简历优化顾问，只能依据已有事实提出修改建议。"),
                ChatMessage(role="user", content=prompt),
            ],
            api_key=resolution.credential,
        )
        payload = self._parse_json_object(raw)
        suggestions = tuple(self._parse_suggestion(item, index) for index, item in enumerate(payload.get("suggestions", []), start=1))
        if not suggestions:
            raise ValueError("模型没有返回可执行的简历修改建议，请重新分析")
        return ResumeAnalysisResult(
            original_markdown=original_markdown,
            job_description_text=combined_job,
            job_title=str(payload.get("job_title") or "目标岗位").strip()[:160],
            analysis_summary=str(payload.get("analysis_summary") or "").strip(),
            suggestions=suggestions,
            model_profile_id=resolution.profile.id,
            model_display_name=resolution.profile.display_name,
            provider_key=resolution.profile.provider_key,
            model_id=resolution.profile.model_id,
        )

    def generate(
        self,
        organization_id: UUID,
        original_markdown: str,
        job_description_text: str,
        suggestions: tuple[ResumeSuggestion, ...],
        extra_prompt: str,
        model_profile_id: UUID,
    ) -> str:
        """按用户勾选的意见生成可转存简历正文，不写历史。"""

        if not suggestions:
            raise ValueError("请至少选择一条修改意见")
        resolution = self._resolve_model(organization_id, model_profile_id)
        prompt = self._generation_prompt(
            original_markdown,
            job_description_text,
            suggestions,
            extra_prompt,
        )
        result = trace_operation(
            run_name="resume_assistant.generate",
            run_type="chain",
            inputs={"suggestion_count": len(suggestions), "source_characters": len(original_markdown)},
            metadata={"component": "resume_assistant"},
            execute=lambda: self._chat_client.complete(
                resolution.profile,
                resolution.credential_env_name,
                [
                    ChatMessage(role="system", content="你是技术简历编辑器。不得虚构经历、技能、指标、公司或项目。"),
                    ChatMessage(role="user", content=prompt),
                ],
                api_key=resolution.credential,
            ),
            summarize=lambda value: {"output_characters": len(value)},
            tags=("career", "resume-assistant"),
        )
        markdown = self._strip_markdown_fence(result).strip()
        markdown = self._ensure_revision_from_original(
            original_markdown=original_markdown,
            optimized_markdown=markdown,
            generation_callback=lambda: self._chat_client.complete(
                resolution.profile,
                resolution.credential_env_name,
                [
                    ChatMessage(
                        role="system",
                        content=(
                            "你是技术简历编辑器。不得虚构经历、技能、公司或项目。"
                            "若上一版偏离原文，请在原文骨架上逐段微调，不得重排事实与时间线。"
                        ),
                    ),
                    ChatMessage(
                        role="user",
                        content=self._build_revision_prompt(
                            original_markdown,
                            job_description_text,
                            suggestions,
                            extra_prompt,
                            strict_mode=True,
                        ),
                    ),
                ],
                api_key=resolution.credential,
            ),
        )
        if len(markdown) < 80:
            raise ValueError("模型返回的优化简历内容过短，请重新生成")
        return markdown

    @classmethod
    def _is_revision_faithful(cls, original_markdown: str, optimized_markdown: str) -> bool:
        """粗粒度判断优化结果是否保留原始简历事实与文本骨架。"""

        if not original_markdown or not optimized_markdown:
            return False
        if len(optimized_markdown) < max(80, int(len(original_markdown) * 0.40)):
            return False
        similarity = SequenceMatcher(None, original_markdown, optimized_markdown).ratio()
        if similarity < 0.25:
            return False
        important = cls._extract_core_facts(original_markdown)
        if not important:
            return True
        matched = sum(1 for token in important if token in optimized_markdown)
        if matched < max(2, int(len(important) * 0.35)):
            return False
        return True

    @staticmethod
    def _extract_core_facts(text: str) -> tuple[str, ...]:
        """提取用于一致性校验的事实锚点（姓名/邮箱/手机号/技能/项目关键词）。"""

        phone_pattern = re.compile(r"(?:1[3-9]\\d{9})")
        email_pattern = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\\.[A-Za-z]{2,}")
        project_pattern = re.compile(r"[\\u4e00-\\u9fa5A-Za-z]{2,18}(?:项目|系统|平台|服务|平台化)")
        skill_pattern = re.compile(r"[A-Za-z][-A-Za-z0-9+/#.]{2,20}")
        raw_tokens = []
        raw_tokens.extend(phone_pattern.findall(text))
        raw_tokens.extend(email_pattern.findall(text))
        raw_tokens.extend(project_pattern.findall(text))
        raw_tokens.extend(skill_pattern.findall(text))

        # 补充去重后的关键短语，优先保留长度较长且非通用词
        stop_words = {
            "项目",
            "经验",
            "负责",
            "熟练",
            "掌握",
            "开发",
            "优化",
            "能力",
            "技术",
            "经验",
            "项目管理",
            "团队",
        }
        normalized: list[str] = []
        for token in raw_tokens:
            t = token.strip()
            if (
                not t
                or len(t) < 3
                or t in stop_words
                or t.lower() in stop_words
                or t.startswith("http")
            ):
                continue
            if t not in normalized:
                normalized.append(t)
            if len(normalized) >= 24:
                break
        return tuple(normalized)

    def _ensure_revision_from_original(
        self,
        *,
        original_markdown: str,
        optimized_markdown: str,
        generation_callback,
    ) -> str:
        """对优化结果做一次原文一致性兜底，防止模型给出非改写文本。"""

        cleaned = self._strip_markdown_fence(optimized_markdown).strip()
        if self._is_revision_faithful(original_markdown, cleaned):
            return cleaned
        retry_raw = self._strip_markdown_fence(str(generation_callback())).strip()
        if self._is_revision_faithful(original_markdown, retry_raw):
            return retry_raw
        raise ValueError("模型生成内容与原始简历偏离较大，请重新生成")

    def _build_revision_prompt(
        self,
        original: str,
        job: str,
        suggestions: tuple[ResumeSuggestion, ...],
        extra: str,
        *,
        strict_mode: bool = False,
    ) -> str:
        """构建重试时更严格的改写提示词，减少生成偏离。"""

        suffix = ""
        if strict_mode:
            suffix = (
                "严格要求：不改变原文段落顺序与时间线，不新增/删减经历、公司、教育经历、项目角色、关键数字；"
                "只在措辞、结构排序、岗位关键词表达上进行微调。"
            )
        return f"""{suffix}
基于原始简历执行以下用户确认的修改意见，输出完整可用于导出的英文/中文文本简历。
硬性规则：
1. 不虚构经历、技能、公司、项目、数字、成果；
2. 保留原始简历的姓名、联系方式、岗位经历、学历和时间线；
3. 修改必须可从“原始简历 + 岗位要求 + 用户意见”中直接推导；
4. 返回文本内容仅包含可直接导出的正文，不允许附加 JSON 或分析说明。

原始简历：
{original}

岗位要求：
{job}

确认意见：
{json.dumps([asdict(item) for item in suggestions], ensure_ascii=False)}

用户额外要求：
{extra.strip() or '无'}"""

    @classmethod
    def render_docx_from_text(cls, text: str, target: Path) -> None:
        """对外提供的简历文档导出能力，供路由下载兜底。"""

        cls._write_docx_file(text, target)

    @classmethod
    def render_pdf_from_text(cls, text: str, target: Path) -> None:
        """对外提供的简历 PDF 兜底导出能力。"""

        cls._write_pdf_from_text(text, target)

    def save_record(
        self,
        *,
        organization_id: UUID,
        owner_id: UUID,
        creator_name: str,
        source_descriptor: AttachmentDescriptor,
        original_preview_markdown: str,
        job_title: str,
        job_description_text: str,
        extra_prompt: str,
        suggestions: tuple[ResumeSuggestion, ...],
        optimized_markdown: str,
        model_profile_id: UUID,
    ) -> ResumeOptimizationRecord:
        """用户点击保存后才复制原件并写入不可变历史。"""

        profile = self._model_repository.get_profile(organization_id, model_profile_id)
        if profile is None:
            raise LookupError("保存记录使用的模型档案不存在")
        record_id = uuid4()
        target_directory = self._storage_root / str(organization_id) / str(record_id)
        target_directory.mkdir(parents=True, exist_ok=False)
        source_name = Path(source_descriptor.original_filename).name
        original_path = target_directory / source_name
        original_preview_pdf_path = target_directory / "original_preview.pdf"
        optimized_text_path = target_directory / "optimized_resume.txt"
        optimized_docx_path = target_directory / "optimized_resume.docx"
        optimized_pdf_path = target_directory / "optimized_resume.pdf"
        shutil.copyfile(source_descriptor.temporary_path, original_path)
        self._write_text_file(optimized_text_path, optimized_markdown)
        self._write_docx_file(optimized_markdown, optimized_docx_path)
        self._write_preview_pdf(source_descriptor, original_preview_pdf_path, original_preview_markdown)
        self._write_pdf_from_text(optimized_markdown, optimized_pdf_path)

        record = ResumeOptimizationRecord(
            id=record_id,
            organization_id=organization_id,
            owner_id=owner_id,
            creator_name=creator_name.strip() or "未命名用户",
            job_title=job_title.strip()[:160] or "目标岗位",
            model_profile_id=profile.id,
            model_display_name=profile.display_name,
            provider_key=profile.provider_key,
            model_id=profile.model_id,
            source_filename=source_name,
            source_media_type=source_descriptor.media_type,
            original_file_path=str(original_path),
            original_preview_pdf_path=str(original_preview_pdf_path),
            original_preview_markdown=original_preview_markdown,
            job_description_text=job_description_text,
            extra_prompt=extra_prompt.strip() or None,
            suggestions=suggestions,
            optimized_markdown=optimized_markdown,
            optimized_text_path=str(optimized_text_path),
            optimized_docx_path=str(optimized_docx_path),
            optimized_pdf_path=str(optimized_pdf_path),
            created_at=datetime.now(timezone.utc),
        )
        return self._repository.create(record)

    def _write_text_file(self, target: Path, content: str) -> None:
        """把优化后的文本写入本地文件，失败时直接抛出。"""

        try:
            target.write_text((content or "").strip() + "\n", encoding="utf-8")
        except OSError as exc:
            raise OSError("优化文本文件写入失败") from exc

    def _write_docx_file(self, optimized_text: str, target: Path) -> None:
        """优先使用 python-docx，不可用时降级到最小 docx 容器。

        降级容器仅保证可打开性；有 python-docx 时按简历常见结构尽量保留标题层级。
        """

        try:
            from docx import Document  # type: ignore
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # type: ignore
            from docx.shared import Pt  # type: ignore

            document = Document()
            lines = optimized_text.splitlines()
            for line in lines:
                text = (line or "").strip()
                if not text:
                    continue
                if re.match(r"^\s*#{1,6}\s+", text):
                    level = len(re.match(r"^\s*(#{1,6})\s+", text).group(1))  # type: ignore
                    heading = text.split("#", level)[-1].strip()
                    document.add_heading(heading, level=level)
                elif re.match(r"^\s*[\-+*]\s+", text):
                    document.add_paragraph(text.lstrip(" -*+\t"), style="List Bullet")
                elif re.match(r"^\s*\d+\.\s+", text):
                    document.add_paragraph(text.lstrip(), style="List Number")
                else:
                    paragraph = document.add_paragraph(text)
                    if "：" in text and paragraph.runs:
                        paragraph.runs[0].bold = False
                        paragraph.runs[0].font.size = Pt(11)
                    if len(lines) == 1:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            document.save(str(target))
            return
        except Exception:
            self._write_docx_fallback(optimized_text, target)

    @staticmethod
    def _write_docx_fallback(optimized_text: str, target: Path) -> None:
        """无外部依赖时生成可打开的最小 docx 容器。"""

        lines = optimized_text.splitlines() or [""]
        paragraph_xml_parts: list[str] = []
        for line in lines:
            if line:
                safe = xml_escape(line)
                paragraph_xml_parts.append(
                    f'<w:p><w:r><w:t xml:space="preserve">{safe}</w:t></w:r></w:p>',
                )
            else:
                paragraph_xml_parts.append("<w:p/>")
        paragraph_xml = "".join(paragraph_xml_parts)

        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            "<w:document xmlns:mc=\"http://schemas.openxmlformats.org/markup-compatibility/2006\" "
            'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:body>"
            f"{paragraph_xml}"
            "<w:sectPr>"
            '<w:pgSz w:w="11906" w:h="16838"/>'
            '<w:pgMar w:top="1440" w:right="1080" w:bottom="1440" w:left="1080" w:header="708" w:footer="708" w:gutter="0"/>'
            "</w:sectPr>"
            "</w:body></w:document>"
        )

        content_types_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            "<Types xmlns=\"http://schemas.openxmlformats.org/package/2006/content-types\">"
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
            '<Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>'
            '<Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>'
            '<Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>'
            '<Override PartName="/word/_rels/document.xml.rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            "</Types>"
        )

        rels_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            "<Relationships xmlns=\"http://schemas.openxmlformats.org/package/2006/relationships\">"
            '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>'
            "</Relationships>"
        )

        package_rels_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            "<Relationships xmlns=\"http://schemas.openxmlformats.org/package/2006/relationships\">"
            '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
            '<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>'
            '<Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>'
            "</Relationships>"
        )

        styles_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:styles xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
            'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )

        core_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" '
            'xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" '
            'xmlns:dcmitype="http://purl.org/dc/dcmitype/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">'
            "<dc:title>optimized_resume</dc:title>"
            "<dc:creator>Career Assistant</dc:creator>"
            "<cp:lastModifiedBy>Career Assistant</cp:lastModifiedBy>"
            f"<dcterms:created xsi:type=\"dcterms:W3CDTF\">{datetime.now(timezone.utc).strftime('%Y-%m-%dT%H:%M:%SZ')}</dcterms:created>"
            "</cp:coreProperties>"
        )

        app_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" '
            'xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">'
            "<Application>Career Assistant</Application>"
            "<Pages>1</Pages>"
            "</Properties>"
        )

        try:
            with zipfile.ZipFile(target, "w", compression=zipfile.ZIP_DEFLATED) as zip_file:
                zip_file.writestr("[Content_Types].xml", content_types_xml)
                zip_file.writestr("_rels/.rels", package_rels_xml)
                zip_file.writestr("word/_rels/document.xml.rels", rels_xml)
                zip_file.writestr("word/document.xml", document_xml)
                zip_file.writestr("word/styles.xml", styles_xml)
                zip_file.writestr("docProps/core.xml", core_xml)
                zip_file.writestr("docProps/app.xml", app_xml)
        except OSError as exc:
            raise OSError("优化简历 DOCX 写入失败") from exc

    def _write_preview_pdf(
        self,
        source_descriptor: AttachmentDescriptor,
        original_preview_pdf_path: Path,
        fallback_text: str,
    ) -> None:
        """尝试保留原始简历样式；失败则用文本 PDF 兜底。"""

        if source_descriptor.temporary_path.suffix.lower() == ".pdf":
            shutil.copyfile(source_descriptor.temporary_path, original_preview_pdf_path)
            return
        if self._legacy_office_converter is not None:
            try:
                converted = self._legacy_office_converter.convert_to_pdf(source_descriptor)
                shutil.copyfile(converted.temporary_path, original_preview_pdf_path)
                return
            except LegacyOfficeConversionError:
                pass
        self._write_pdf_from_text(fallback_text, original_preview_pdf_path)

    @staticmethod
    def _write_pdf_from_text(text: str, output: Path) -> None:
        """无 PDF 能力时回退到可读的文本 PDF，保证下载链路可验证。"""

        lines = [line.strip() for line in (text or "").splitlines() if line.strip()]
        if not lines:
            lines = ["（该简历暂未生成可预览文本）"]

        def _escape_pdf_string(value: str) -> str:
            safe = value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
            return safe.replace("\r", "")

        def _build_text_block() -> str:
            rows = ["BT", "/F1 11 Tf", "72 760 Td"]
            for line in lines[:120]:
                rows.append(f"({_escape_pdf_string(line)}) Tj")
                rows.append("0 -14 Td")
            rows.append("ET")
            return "\n".join(rows)

        content = _build_text_block().encode("latin1", errors="replace")
        content_stream = b"".join(
            [
                b"<< /Length ",
                str(len(content)).encode("ascii"),
                b" >>\nstream\n",
                content,
                b"\nendstream",
            ],
        )

        object_defs: list[bytes] = [
            b"<< /Type /Catalog /Pages 2 0 R >>",
            b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
            b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica /Encoding /WinAnsiEncoding >>",
            content_stream,
        ]

        objects: list[bytes] = [b"%PDF-1.4\n"]
        offsets: list[int] = []
        for index, value in enumerate(object_defs, start=1):
            offsets.append(len(b"".join(objects)))
            objects.append(f"{index} 0 obj\n".encode("ascii"))
            objects.append(value)
            objects.append(b"\nendobj\n")

        xref_offset = len(b"".join(objects))
        trailer = ["xref\n", "0 6\n", "0000000000 65535 f \n"]
        for offset in offsets:
            trailer.append(f"{offset:010d} 00000 n \n")
        trailer.extend(
            [
                "trailer << /Size 6 /Root 1 0 R >>\n",
                "startxref\n",
                f"{xref_offset}\n",
                "%%EOF\n",
            ],
        )
        trailer_bytes = "".join(trailer).encode("ascii")
        pdf_bytes = b"".join(objects) + trailer_bytes
        try:
            output.write_bytes(pdf_bytes)
        except OSError as exc:
            raise OSError("PDF 文件写入失败") from exc

    def _resolve_model(self, organization_id: UUID, profile_id: UUID | None) -> ModelResolution:
        if profile_id is None:
            profiles = self._model_repository.list_profiles(organization_id, include_disabled=False)
            profile = next(
                (
                    item
                    for item in profiles
                    if "deepseekv4flash" in re.sub(
                        r"[^a-z0-9]",
                        "",
                        f"{item.profile_key}{item.display_name}{item.model_id}".lower(),
                    )
                ),
                None,
            )
            if profile is None:
                raise LookupError("未找到默认模型 deepseek-v4-flash，请先在模型库配置")
            profile_id = profile.id
        resolution = self._model_gateway.resolve(
            organization_id,
            ModelSelectionRequest(
                mode=ModelSelectionMode.SPECIFIC_PROFILE,
                profile_id=profile_id,
                required_capabilities=frozenset({ModelCapability.TEXT}),
            ),
        )
        if resolution.readiness is not ModelReadiness.READY:
            raise ValueError("所选模型尚不可调用，请检查模型连接")
        return resolution

    @staticmethod
    def _analysis_prompt(original: str, job: str, extra: str) -> str:
        return f"""分析下面的原始简历与岗位要求，返回严格 JSON，不要 Markdown 代码块。
JSON 格式：
{{"job_title":"岗位名称","analysis_summary":"整体匹配结论","suggestions":[{{"id":"s1","category":"结构|表达|匹配|关键词","title":"短标题","rationale":"为什么要改","original_evidence":"简历中的依据","job_evidence":"岗位中的依据","proposed_change":"具体怎么改","priority":"high|medium|low"}}]}}
规则：只依据原始简历中的事实；禁止虚构经历、技能、指标、公司、项目；建议必须能由用户独立勾选；合并重复意见。

原始简历：
{original}

岗位要求：
{job}

用户额外要求：
{extra.strip() or '无'}"""

    @staticmethod
    def _generation_prompt(original: str, job: str, suggestions: tuple[ResumeSuggestion, ...], extra: str) -> str:
        return f"""基于原始简历，执行用户已确认的修改意见，输出一份可直接用于导出 Word/PDF 的完整简历正文文本。
硬性规则：不得新增原文不存在的经历、公司、项目、技能、数字或成果；保留联系方式与事实；只优化结构、顺序、措辞与岗位相关表达；不要输出 Markdown 语法、不要输出分析过程。

原始简历：
{original}

岗位要求：
{job}

已确认意见：
{json.dumps([asdict(item) for item in suggestions], ensure_ascii=False)}

用户额外要求：
{extra.strip() or '无'}"""

    @staticmethod
    def _parse_json_object(raw: str) -> dict[str, object]:
        cleaned = ResumeOptimizationService._strip_markdown_fence(raw).strip()
        start, end = cleaned.find("{"), cleaned.rfind("}")
        if start < 0 or end <= start:
            raise ValueError("模型返回内容不是有效 JSON")
        try:
            payload = json.loads(cleaned[start : end + 1])
        except json.JSONDecodeError as exc:
            raise ValueError("模型返回的分析结构无法解析，请重新分析") from exc
        if not isinstance(payload, dict):
            raise ValueError("模型返回的分析结构无效")
        return payload

    @staticmethod
    def _parse_suggestion(item: object, index: int) -> ResumeSuggestion:
        if not isinstance(item, dict):
            raise ValueError("模型返回了无效的修改意见")
        value = lambda key, fallback="": str(item.get(key) or fallback).strip()
        return ResumeSuggestion(
            id=value("id", f"s{index}"),
            category=value("category", "表达"),
            title=value("title", f"修改建议 {index}"),
            rationale=value("rationale"),
            original_evidence=value("original_evidence"),
            job_evidence=value("job_evidence"),
            proposed_change=value("proposed_change"),
            priority=value("priority", "medium"),
        )

    @staticmethod
    def _strip_markdown_fence(value: str) -> str:
        stripped = value.strip()
        if stripped.startswith("```"):
            stripped = re.sub(r"^```(?:json|markdown|md)?\s*", "", stripped, flags=re.I)
            stripped = re.sub(r"\s*```$", "", stripped)
        return stripped
